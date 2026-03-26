"""
多模态文档处理模块
支持 PDF、Word、Excel 中的表格提取
支持图片 OCR 和 Vision LLM 内容理解
"""
from typing import List, Dict, Any, Optional, Union
from pathlib import Path
import re
import json
import io
import os
import tempfile
import logging
import asyncio

from langchain_core.documents import Document

from config.settings import get_settings

logger = logging.getLogger(__name__)

# ============================================================
# 表格提取器
# ============================================================

class TableExtractor:
    """表格提取器：支持从 PDF、Word、Excel 中提取表格"""

    def __init__(self):
        self.settings = get_settings()

    def extract_from_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        """从 PDF 中提取表格"""
        tables = []
        try:
            import fitz
            doc = fitz.open(file_path)
            for page_num, page in enumerate(doc):
                tables.extend(self._extract_tables_from_page(page, page_num + 1))
            doc.close()
        except ImportError:
            logger.warning("请安装 pymupdf 来支持 PDF 表格提取: pip install pymupdf")
        except Exception as e:
            logger.warning(f"PDF 表格提取失败: {e}")
        return tables

    def _extract_tables_from_page(self, page, page_num: int) -> List[Dict[str, Any]]:
        """从单页 PDF 中提取表格（文本块分析简化版）"""
        tables = []
        try:
            blocks = page.get_text("blocks")
            # 简化实现：识别等宽排列的文本块作为表格
            # 完整实现可使用 camelot、tabula 等库
            pass
        except Exception as e:
            logger.warning(f"页面表格提取失败: {e}")
        return tables

    def extract_from_docx(self, file_path: str) -> List[Dict[str, Any]]:
        """从 Word 文档中提取表格"""
        tables = []
        try:
            from docx import Document
            doc = Document(file_path)
            for i, table in enumerate(doc.tables):
                table_data = {"table_index": i, "rows": [], "headers": []}
                if table.rows:
                    headers = [cell.text.strip() for cell in table.rows[0].cells]
                    table_data["headers"] = headers
                    for row in table.rows[1:]:
                        row_data = [cell.text.strip() for cell in row.cells]
                        table_data["rows"].append(row_data)
                tables.append(table_data)
        except ImportError:
            logger.warning("请安装 python-docx 来支持 Word 表格提取: pip install python-docx")
        except Exception as e:
            logger.warning(f"Word 表格提取失败: {e}")
        return tables

    def extract_from_excel(self, file_path: str) -> List[Dict[str, Any]]:
        """从 Excel 文件中提取表格"""
        tables = []
        try:
            import openpyxl
            wb = openpyxl.load_workbook(file_path, data_only=True)
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                table_data = {"sheet_name": sheet_name, "headers": [], "rows": []}
                headers = [str(cell.value) for cell in sheet[1] if cell.value is not None]
                table_data["headers"] = headers
                for row in sheet.iter_rows(min_row=2):
                    row_data = [str(cell.value or "") for cell in row]
                    if any(row_data):
                        table_data["rows"].append(row_data)
                tables.append(table_data)
            wb.close()
        except ImportError:
            logger.warning("请安装 openpyxl 来支持 Excel 表格提取: pip install openpyxl")
        except Exception as e:
            logger.warning(f"Excel 表格提取失败: {e}")
        return tables

    def table_to_markdown(self, table: Dict[str, Any]) -> str:
        """将表格转换为 Markdown 格式"""
        if not table.get("rows"):
            return ""
        headers = table.get("headers", [])
        rows = table.get("rows", [])
        lines = []
        if headers:
            lines.append("| " + " | ".join(headers) + " |")
            lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
        for row in rows:
            lines.append("| " + " | ".join(row[:len(headers)]) + " |")
        return "\n".join(lines)


# ============================================================
# 图片提取器（重写：提取真实 bytes）
# ============================================================

class ImageExtractor:
    """
    图片提取器
    从 PDF、DOCX 中提取图片原始数据（bytes）
    """

    def __init__(self, skip_small_px: int = 64):
        self.settings = get_settings()
        self.skip_small_px = skip_small_px

    def extract_from_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        """
        从 PDF 中提取所有图片的 bytes 和元数据

        Returns:
            图片列表，每项包含 bytes、格式、尺寸、页码
        """
        images = []
        seen_xrefs = set()  # 按 xref 去重，同一张图跨页不重复

        try:
            import fitz
            doc = fitz.open(file_path)

            for page_num, page in enumerate(doc):
                image_list = page.get_images(full=True)
                for img_index, img in enumerate(image_list):
                    try:
                        xref = img[0]

                        # 跳过已见过的 xref（同图跨页）
                        if xref in seen_xrefs:
                            continue
                        seen_xrefs.add(xref)

                        base_image = doc.extract_image(xref)
                        ext = base_image["ext"]
                        width = base_image["width"]
                        height = base_image["height"]
                        image_bytes = base_image["image"]

                        # 跳过太小的图片（图标、水印等）
                        if width < self.skip_small_px or height < self.skip_small_px:
                            logger.debug(f"跳过小图片 {width}x{height} (页 {page_num + 1})")
                            continue

                        images.append({
                            "page": page_num + 1,
                            "index": img_index,
                            "xref": xref,
                            "image_type": ext,
                            "width": width,
                            "height": height,
                            "image_bytes": image_bytes,
                            "file_path": file_path,
                        })
                    except Exception as e:
                        logger.warning(f"提取第 {page_num + 1} 页图片失败: {e}")

            doc.close()
            logger.info(f"PDF 图片提取完成: {len(images)} 张去重图片 (来自 {file_path})")

        except ImportError:
            logger.warning("请安装 pymupdf 来支持 PDF 图片提取: pip install pymupdf")
        except Exception as e:
            logger.warning(f"PDF 图片提取失败: {e}")

        return images

    def extract_from_docx(self, file_path: str) -> List[Dict[str, Any]]:
        """
        从 Word 文档中提取所有图片的 bytes 和元数据
        """
        images = []
        try:
            from docx import Document
            import zipfile

            doc = Document(file_path)
            # DOCX 本质是 zip，图片存储在 word/media/
            zip_path = file_path  # .docx 文件本身是 zip

            try:
                with zipfile.ZipFile(zip_path, "r") as zf:
                    media_files = [f for f in zf.namelist() if f.startswith("word/media/")]
                    for idx, media_path in enumerate(media_files):
                        try:
                            image_bytes = zf.read(media_path)
                            # 获取图片尺寸
                            from PIL import Image
                            img_obj = Image.open(io.BytesIO(image_bytes))
                            width, height = img_obj.size
                            img_obj.close()

                            if width < self.skip_small_px or height < self.skip_small_px:
                                continue

                            ext = Path(media_path).suffix.lstrip(".").lower()
                            images.append({
                                "page": 0,  # DOCX 没有页码概念
                                "index": idx,
                                "xref": 0,
                                "image_type": ext,
                                "width": width,
                                "height": height,
                                "image_bytes": image_bytes,
                                "file_path": file_path,
                            })
                        except Exception as e:
                            logger.warning(f"提取 DOCX 图片失败: {e}")
            except zipfile.BadZipFile:
                logger.warning(f"无法打开 DOCX 文件: {file_path}")

            logger.info(f"DOCX 图片提取完成: {len(images)} 张有效图片 (来自 {file_path})")

        except ImportError:
            logger.warning("请安装 python-docx 和 pillow 来支持 Word 图片提取")
        except Exception as e:
            logger.warning(f"DOCX 图片提取失败: {e}")

        return images


# ============================================================
# OCR 处理器
# ============================================================

class OCRProcessor:
    """
    OCR 处理器
    使用 Tesseract 识别图片中的文字
    """

    def __init__(self):
        self.settings = get_settings()

    def process_image_bytes(
        self,
        image_bytes: bytes,
        lang: str = "chi_sim+eng",
        dpi: int = 200,
    ) -> str:
        """
        对图片 bytes 进行 OCR 识别

        Args:
            image_bytes: 图片数据
            lang: 识别语言（chi_sim+eng 中文+英文）
            dpi: OCR 分辨率

        Returns:
            识别出的文本
        """
        try:
            import pytesseract
            from PIL import Image

            image = Image.open(io.BytesIO(image_bytes))
            # 提高分辨率以改善识别效果
            if image.mode not in ("RGB", "L"):
                image = image.convert("RGB")
            text = pytesseract.image_to_string(image, lang=lang)
            image.close()
            return text.strip()

        except ImportError:
            logger.debug("pytesseract 未安装，跳过 OCR")
            return ""
        except Exception as e:
            logger.warning(f"OCR 处理失败: {e}")
            return ""

    def process_image_file(
        self,
        image_path: str,
        lang: str = "chi_sim+eng",
    ) -> str:
        """对图片文件进行 OCR 识别"""
        try:
            with open(image_path, "rb") as f:
                return self.process_image_bytes(f.read(), lang)
        except Exception as e:
            logger.warning(f"OCR 文件处理失败: {e}")
            return ""


# ============================================================
# Vision LLM 图片理解处理器（新增核心组件）
# ============================================================

class VisionProcessor:
    """
    Vision LLM 图片理解处理器
    使用 qwen-vl-plus 等多模态模型理解文档中的图片
    """

    def __init__(self):
        self.settings = get_settings()

    def _create_vision_llm(self):
        """创建 Vision LLM 实例"""
        import os as _os
        from langchain_openai import ChatOpenAI

        # 设置代理
        http_proxy = _os.environ.get("http_proxy") or _os.environ.get("HTTP_PROXY")
        https_proxy = _os.environ.get("https_proxy") or _os.environ.get("HTTPS_PROXY")
        if http_proxy and not _os.environ.get("HTTP_PROXY"):
            _os.environ["HTTP_PROXY"] = http_proxy
        if https_proxy and not _os.environ.get("HTTPS_PROXY"):
            _os.environ["HTTPS_PROXY"] = https_proxy

        return ChatOpenAI(
            model=self.settings.vision_ingestion_model,
            temperature=0.1,
            max_tokens=2048,
            api_key=self.settings.dashscope_api_key,
            base_url=self.settings.dashscope_base_url,
        )

    def _compress_image(
        self,
        image_bytes: bytes,
        max_size_mb: float = 5.0,
        max_dimension: int = 2048,
    ) -> bytes:
        """压缩图片到指定大小和尺寸"""
        try:
            from PIL import Image
            img = Image.open(io.BytesIO(image_bytes))

            # 缩放
            if max(img.size) > max_dimension:
                ratio = max_dimension / max(img.size)
                new_size = (int(img.size[0] * ratio), int(img.size[1] * ratio))
                img = img.resize(new_size, Image.LANCZOS)

            output = io.BytesIO()
            fmt = img.format or "JPEG"
            img.save(output, format=fmt, quality=85)
            compressed = output.getvalue()
            img.close()

            # 如果仍然太大，继续压缩
            if len(compressed) / (1024 * 1024) > max_size_mb:
                img2 = Image.open(io.BytesIO(compressed))
                output2 = io.BytesIO()
                img2.save(output2, format=fmt, quality=60)
                compressed = output2.getvalue()
                img2.close()

            return compressed

        except Exception as e:
            logger.warning(f"图片压缩失败: {e}")
            return image_bytes

    def _encode_image(self, image_bytes: bytes) -> str:
        """将图片 bytes 转为 data URI 格式"""
        try:
            from PIL import Image
            img = Image.open(io.BytesIO(image_bytes))
            fmt = (img.format or "jpeg").lower()
            mime = f"image/{fmt}"
            b64 = __import__("base64").b64encode(image_bytes).decode("utf-8")
            img.close()
            return f"data:{mime};base64,{b64}"
        except Exception:
            # 回退：当作 jpeg
            b64 = __import__("base64").b64encode(image_bytes).decode("utf-8")
            return f"data:image/jpeg;base64,{b64}"

    def _build_vision_content(
        self,
        images: List[Dict[str, Any]],
        prompt: str,
    ) -> List[Dict[str, Any]]:
        """
        构建 Vision LLM 的多模态消息内容

        Args:
            images: 图片数据列表
            prompt: 提示词

        Returns:
            LangChain 格式的 content 列表
        """
        content = []
        for img in images:
            # 压缩并编码
            compressed = self._compress_image(img["image_bytes"])
            data_uri = self._encode_image(compressed)
            content.append({
                "type": "image_url",
                "image_url": {"url": data_uri},
            })

        content.append({"type": "text", "text": prompt})
        return content

    async def understand_images_async(
        self,
        images: List[Dict[str, Any]],
        prompt: str = None,
    ) -> List[Dict[str, Any]]:
        """
        异步理解多张图片（逐张调用 Vision LLM，避免单次上下文过长）

        Args:
            images: 图片列表（包含 image_bytes）
            prompt: 理解提示词

        Returns:
            每张图片的理解结果列表
        """
        if not images:
            return []

        settings = self.settings
        if not getattr(settings, "vision_ingestion_enabled", True):
            logger.info("Vision 入库已禁用，跳过图片理解")
            return [{"error": "vision_disabled"}] * len(images)

        prompt = prompt or settings.vision_ingestion_prompt
        llm = self._create_vision_llm()

        from langchain_core.messages import HumanMessage
        results = []

        for i, img in enumerate(images):
            try:
                logger.info(
                    f"[Vision] 理解图片 {i + 1}/{len(images)}: "
                    f"{img.get('width', 0)}x{img.get('height', 0)}, "
                    f"页 {img.get('page', '?')}"
                )

                content = self._build_vision_content([img], prompt)
                message = HumanMessage(content=content)

                response = await llm.ainvoke([message])
                description = response.content

                results.append({
                    "success": True,
                    "description": description,
                    "page": img.get("page"),
                    "index": img.get("index"),
                    "width": img.get("width"),
                    "height": img.get("height"),
                })

                logger.info(f"[Vision] 图片 {i + 1} 理解成功，长度: {len(description)}")

            except Exception as e:
                logger.warning(f"[Vision] 图片 {i + 1} 理解失败: {e}")
                results.append({
                    "success": False,
                    "error": str(e),
                    "page": img.get("page"),
                    "index": img.get("index"),
                })

        return results

    def understand_images(
        self,
        images: List[Dict[str, Any]],
        prompt: str = None,
    ) -> List[Dict[str, Any]]:
        """
        同步版本：在新事件循环中执行异步理解
        """
        try:
            loop = asyncio.get_running_loop()
        except RuntimeError:
            # 没有运行中的循环，直接创建
            return asyncio.run(self.understand_images_async(images, prompt))

        # 已有运行中的循环，用 ThreadPoolExecutor 包装
        import concurrent.futures
        with concurrent.futures.ThreadPoolExecutor() as pool:
            future = pool.submit(
                asyncio.run,
                self.understand_images_async(images, prompt)
            )
            return future.result()

    def process_images_for_document(
        self,
        file_path: str,
        images: List[Dict[str, Any]],
        document_title: str = "",
    ) -> List[Dict[str, Any]]:
        """
        处理文档中的所有图片，返回理解结果

        Args:
            file_path: 文档路径
            images: 从文档中提取的图片列表
            document_title: 文档标题（用于提示词上下文）

        Returns:
            每张图片的理解结果
        """
        settings = self.settings

        if not images:
            return []

        # 限制数量
        max_images = settings.vision_ingestion_max_images_per_doc
        if len(images) > max_images:
            logger.warning(
                f"文档 {file_path} 包含 {len(images)} 张图片，"
                f"超过限制 {max_images}，只处理前 {max_images} 张"
            )
            images = images[:max_images]

        # 构建上下文感知提示词
        base_prompt = settings.vision_ingestion_prompt
        if document_title:
            prompt = (
                f"【文档背景】这份文档的标题是「{document_title}」。\n\n"
                f"{base_prompt}"
            )
        else:
            prompt = base_prompt

        logger.info(
            f"[Vision] 开始处理文档 {Path(file_path).name}，"
            f"共 {len(images)} 张图片"
        )

        return self.understand_images(images, prompt=prompt)


# ============================================================
# 多模态文档处理器（重写：接入 Vision）
# ============================================================

class MultimodalDocumentProcessor:
    """
    多模态文档处理器
    整合表格提取、图片提取（Vision LLM）、OCR 功能
    """

    def __init__(self):
        self.settings = get_settings()
        self.table_extractor = TableExtractor()
        self.image_extractor = ImageExtractor(
            skip_small_px=self.settings.vision_ingestion_skip_small
        )
        self.ocr_processor = OCRProcessor()
        self.vision_processor = VisionProcessor()

    def _extract_images_from_document(
        self,
        file_path: str,
    ) -> List[Dict[str, Any]]:
        """
        从文档中提取所有图片

        支持格式: .pdf, .docx, .doc
        """
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix == ".pdf":
            return self.image_extractor.extract_from_pdf(file_path)
        elif suffix in [".docx", ".doc"]:
            return self.image_extractor.extract_from_docx(file_path)
        else:
            return []

    def _build_image_description_text(
        self,
        results: List[Dict[str, Any]],
        file_name: str,
    ) -> str:
        """
        将 Vision LLM 的图片理解结果格式化为文本片段
        追加到文档内容中
        """
        if not results:
            return ""

        descriptions = []
        for i, r in enumerate(results):
            page_info = f"第{r.get('page', '?')}页" if r.get("page", 0) > 0 else "文档内"
            img_info = f"[图片 {i + 1} ({page_info}, {r.get('width', 0)}x{r.get('height', 0)}像素)]"

            if r.get("success") and r.get("description"):
                descriptions.append(f"{img_info}\n{r['description']}")
            elif r.get("error"):
                descriptions.append(f"{img_info}\n[图片理解失败: {r.get('error')}]")
            else:
                descriptions.append(f"{img_info}\n[图片理解跳过]")

        header = f"\n\n## 文档图片内容（共 {len(results)} 张图片）\n\n"
        return header + "\n\n".join(descriptions)

    def enhance_document(
        self,
        document: Document,
    ) -> Document:
        """
        增强单个文档：
        1. 从文档中提取图片 bytes
        2. 调用 Vision LLM 理解图片
        3. 将图片描述追加到 page_content
        4. 更新 metadata
        """
        file_path = document.metadata.get("file_path") or document.metadata.get("source_file")
        if not file_path or not Path(file_path).exists():
            return document

        suffix = Path(file_path).suffix.lower()
        if suffix not in [".pdf", ".docx", ".doc"]:
            return document

        document_title = document.metadata.get("document_title", "")
        file_name = Path(file_path).name

        # 提取图片
        images = self._extract_images_from_document(file_path)

        if not images:
            logger.debug(f"文档 {file_name} 不含图片")
            return document

        logger.info(
            f"[Multimodal] 增强文档: {file_name}, "
            f"发现 {len(images)} 张图片"
        )

        # Vision LLM 理解
        vision_results = self.vision_processor.process_images_for_document(
            file_path=file_path,
            images=images,
            document_title=document_title,
        )

        # 将图片描述追加到 page_content
        if vision_results:
            description_text = self._build_image_description_text(
                vision_results,
                file_name,
            )
            document.page_content += description_text

            # 更新 metadata
            successful = sum(1 for r in vision_results if r.get("success"))
            document.metadata["has_images"] = True
            document.metadata["images_count"] = len(images)
            document.metadata["images_understood"] = successful
            document.metadata["images_as_text"] = True

        return document

    def enhance_documents(
        self,
        documents: List[Document],
    ) -> List[Document]:
        """
        批量增强文档列表：
        对每个文档执行图片提取和 Vision LLM 理解，
        图片描述广播到该文档的所有 chunk

        Args:
            documents: 原始文档列表

        Returns:
            增强后的文档列表
        """
        if not self.settings.vision_ingestion_enabled:
            print(f"      Vision 入库未启用，跳过图片理解")
            return documents

        enhanced = []
        total_images = 0
        total_docs_with_images = 0

        # 按 file_path 去重：同一文件只处理一次图片
        processed_files = set()

        for i, doc in enumerate(documents):
            file_path = doc.metadata.get("file_path") or doc.metadata.get("source_file")
            file_name = doc.metadata.get("file_name", "?")

            # 非图片类文档不处理
            if not file_path or not Path(file_path).exists():
                enhanced.append(doc)
                continue

            suffix = Path(file_path).suffix.lower()
            if suffix not in [".pdf", ".docx", ".doc"]:
                enhanced.append(doc)
                continue

            # 同一文件只处理一次图片
            already_done = file_path in processed_files
            if already_done:
                enhanced.append(doc)
                continue
            processed_files.add(file_path)

            # 增强单个文档（只对第一个页面文档执行图片处理）
            enhanced_doc = self.enhance_document(doc)
            enhanced.append(enhanced_doc)

            # 统计
            if enhanced_doc.metadata.get("has_images"):
                total_docs_with_images += 1
                total_images += enhanced_doc.metadata.get("images_count", 0)
                print(
                    f"      [Vision] 文档 {i + 1}/{len(documents)}: {file_name} | "
                    f"图片 {enhanced_doc.metadata.get('images_count', 0)} 张"
                )
            elif (i + 1) % 5 == 0:
                print(f"      [Vision] 文档 {i + 1}/{len(documents)}: {file_name} (无图片)")

        print(
            f"      [Vision] 图片理解完成: "
            f"{total_docs_with_images} 个文档含图片, "
            f"共处理 {total_images} 张图片"
        )
        return enhanced


# ============================================================
# 全局实例
# ============================================================

_multimodal_processor: Optional[MultimodalDocumentProcessor] = None


def get_multimodal_processor() -> MultimodalDocumentProcessor:
    """获取多模态文档处理器实例"""
    global _multimodal_processor
    if _multimodal_processor is None:
        _multimodal_processor = MultimodalDocumentProcessor()
    return _multimodal_processor
